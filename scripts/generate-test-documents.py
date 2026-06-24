"""
Generate test fixtures for the document intelligence pipeline.
Creates 9 different document types in /tmp/test-docs/.

Run: python3 scripts/generate-test-documents.py
"""
import os, sys

OUT_DIR = '/tmp/test-docs'
os.makedirs(OUT_DIR, exist_ok=True)

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
except ImportError:
    print("ERROR: reportlab not installed. pip install reportlab")
    sys.exit(1)


def make_searchable_pdf():
    """PDF with proper embedded text — should extract via pdf_text only."""
    path = os.path.join(OUT_DIR, 'searchable_estimate.pdf')
    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(100, 750, "ABC Roofing Company")
    c.drawString(100, 730, "Estimate - Johnson Residence")
    c.drawString(100, 710, "Insured: Sarah Johnson")
    c.drawString(100, 690, "Claim Number: CLM-2024-1234")
    c.drawString(100, 670, "Property: 142 Maple Street, Springfield, IL 62701")
    c.drawString(100, 650, "Insurance Company: State Farm")
    c.drawString(100, 630, "Date of Loss: 03/15/2024")
    c.drawString(100, 610, "Adjuster: Mike Thompson")
    c.drawString(100, 590, "Adjuster Phone: (555) 123-4567")
    c.drawString(100, 570, "Adjuster Email: mthompson@statefarm.com")
    c.drawString(100, 540, "1. Tear off existing shingles - 28 SQ - $4,200.00")
    c.drawString(100, 520, "2. Install GAF Timberline HDZ - 28 SQ - $8,400.00")
    c.drawString(100, 500, "3. Tiger Paw underlayment - 28 SQ - $1,400.00")
    c.drawString(100, 480, "4. Ice and water shield - 6 LF - $480.00")
    c.drawString(100, 460, "5. Drip edge - 180 LF - $720.00")
    c.drawString(100, 430, "RCV: $15,200.00")
    c.drawString(100, 410, "Deductible: $1,000.00")
    c.drawString(100, 390, "ACV: $13,200.00")
    c.drawString(100, 370, "Depreciation: $1,000.00")
    c.save()
    return path


def make_scanned_pdf():
    """PDF with NO text — just shapes. Should trigger OCR fallback."""
    path = os.path.join(OUT_DIR, 'scanned_inspection.pdf')
    c = canvas.Canvas(path, pagesize=letter)
    # Just draw shapes — no text. Simulates a scanned document.
    c.rect(50, 50, 500, 700, fill=0, stroke=1)
    c.rect(60, 60, 480, 680, fill=0, stroke=1)
    c.line(60, 600, 540, 600)  # horizontal divider
    c.line(60, 500, 540, 500)
    c.line(60, 400, 540, 400)
    c.line(60, 300, 540, 300)
    c.line(60, 200, 540, 200)
    # Some "stamp" looking circles
    c.circle(450, 100, 30, fill=0, stroke=1)
    c.circle(450, 100, 25, fill=0, stroke=1)
    c.save()
    return path


def make_mixed_pdf():
    """PDF with text on page 1, blank/scanned page 2."""
    path = os.path.join(OUT_DIR, 'mixed_estimate.pdf')
    c = canvas.Canvas(path, pagesize=letter)
    # Page 1 — has text
    c.drawString(100, 750, "State Farm Insurance")
    c.drawString(100, 730, "Estimate Resolution")
    c.drawString(100, 710, "Insured: David Chen")
    c.drawString(100, 690, "Claim Number: CLM-2024-5678")
    c.drawString(100, 670, "Property: 215 Cedar Court, Lakeside")
    c.drawString(100, 650, "Date of Loss: 02/20/2024")
    c.drawString(100, 620, "1. Remove and dispose old shingles - 30 SQ")
    c.drawString(100, 600, "2. Install architectural shingles - 30 SQ")
    c.drawString(100, 580, "3. Replace roof decking - 5 SF")
    c.drawString(100, 550, "RCV: $18,500.00")
    c.drawString(100, 530, "Deductible: $1,500.00")
    c.drawString(100, 510, "ACV: $16,000.00")
    c.showPage()
    # Page 2 — blank (simulates scanned attachment)
    c.rect(50, 50, 500, 700, fill=0, stroke=1)
    c.save()
    return path


def make_carrier_letter_pdf():
    """Carrier letter — text-based, has 'Dear Insured', 'Sincerely', etc."""
    path = os.path.join(OUT_DIR, 'carrier_letter.pdf')
    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(100, 750, "State Farm Insurance Company")
    c.drawString(100, 730, "P.O. Box 1234")
    c.drawString(100, 710, "Bloomington, IL 61701")
    c.drawString(100, 670, "March 22, 2024")
    c.drawString(100, 640, "Dear Insured,")
    c.drawString(100, 610, "We have completed our review of your claim CLM-2024-1234.")
    c.drawString(100, 590, "After careful consideration, we have approved the estimate")
    c.drawString(100, 570, "in the amount of $15,200.00 RCV.")
    c.drawString(100, 540, "Policy Number: SF-987654321")
    c.drawString(100, 520, "Date of Loss: 03/15/2024")
    c.drawString(100, 490, "If you have any questions, please contact your adjuster:")
    c.drawString(100, 470, "Mike Thompson (555) 123-4567")
    c.drawString(100, 440, "Sincerely,")
    c.drawString(100, 410, "Claims Department")
    c.drawString(100, 380, "State Farm Insurance")
    c.save()
    return path


def make_inspection_report_pdf():
    """Roof inspection report."""
    path = os.path.join(OUT_DIR, 'inspection_report.pdf')
    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(100, 750, "Roof Inspection Report")
    c.drawString(100, 730, "Property: 88 Oak Avenue, Riverside")
    c.drawString(100, 710, "Customer: Carlos Martinez")
    c.drawString(100, 690, "Inspector: James Wilson")
    c.drawString(100, 670, "Date: 04/10/2024")
    c.drawString(100, 640, "Findings:")
    c.drawString(100, 620, "- Hail damage to north slope")
    c.drawString(100, 600, "- 12 shingles missing on east slope")
    c.drawString(100, 580, "- Flashing damaged around chimney")
    c.drawString(100, 560, "- Granule loss visible on south slope")
    c.drawString(100, 530, "Recommendation: Full roof replacement")
    c.drawString(100, 510, "Estimated squares: 28 SQ")
    c.drawString(100, 480, "Photos attached: 8")
    c.save()
    return path


def make_text_file():
    """Plain text file — should be read directly, fileType='other'."""
    path = os.path.join(OUT_DIR, 'customer_notes.txt')
    content = """Customer Notes - Johnson Reroof Project
=====================================

Date: 04/12/2024
Customer: Sarah Johnson
Phone: (555) 200-1001
Email: sarah.j@gmail.com
Address: 142 Maple Street, Springfield, IL 62701

Conversation summary:
- Customer wants to upgrade from 3-tab to architectural shingles
- Color preference: Charcoal (GAF Timberline HDZ)
- Asked about warranty — explained GAF Golden Pledge
- Quote provided: $18,500 for full replacement
- Insurance claim filed: CLM-2024-1234 with State Farm
- Adjuster Mike Thompson scheduled for 04/15/2024

Next steps:
1. Wait for insurance approval
2. Schedule tear-off crew (Ramirez Roofing)
3. Order materials from ABC Supply
"""
    with open(path, 'w') as f:
        f.write(content)
    return path


def make_csv_file():
    """CSV file — material list."""
    path = os.path.join(OUT_DIR, 'material_list.csv')
    content = """SKU,Name,Category,Unit,UnitCost
GAF-THDZ-CHAR,Timberline HDZ Charcoal,Shingles,SQ,105.00
GAF-TP-200,Tiger Paw Underlayment,Underlayment,ROLL,185.00
GAF-STARTER,Starter Strip,Starter,LF,1.85
GAF-TIMBERTEX,Timbertex Hip & Ridge,Hip & Ridge,LF,4.25
DE-DripEdge10,Drip Edge 10ft,Drip Edge,LF,2.10
ABC-Nails15,Roofing Nails 15lb,Fasteners,BOX,42.00
GAF-IWS,Ice & Water Shield,Underlayment,LF,2.85
"""
    with open(path, 'w') as f:
        f.write(content)
    return path


def make_image_files():
    """Create a PNG and JPG with text via reportlab + PIL fallback."""
    paths = []
    # Create PNG via reportlab — draw a "document-like" image
    png_path = os.path.join(OUT_DIR, 'photo_damage.png')
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (800, 600), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        # Draw a simple "roof with damage" representation
        draw.polygon([(100, 400), (400, 100), (700, 400)], fill=(139, 69, 19), outline=(0, 0, 0))
        # Damage marks
        draw.ellipse([300, 280, 340, 320], fill=(0, 0, 0))
        draw.ellipse([420, 250, 460, 290], fill=(0, 0, 0))
        draw.ellipse([500, 320, 540, 360], fill=(0, 0, 0))
        # Labels
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
        except:
            font = ImageFont.load_default()
        draw.text((50, 30), "Roof Damage - 142 Maple Street", fill=(0, 0, 0), font=font)
        draw.text((50, 50), "Hail impact marks visible", fill=(0, 0, 0), font=font)
        draw.text((50, 500), "Inspected: 04/10/2024", fill=(0, 0, 0), font=font)
        img.save(png_path)
        paths.append(png_path)
    except ImportError:
        # Fallback: reportlab-only
        from reportlab.pdfgen import canvas as rl_canvas
        # Skip image generation if PIL not available
        pass

    # Create a simple JPG too
    jpg_path = os.path.join(OUT_DIR, 'photo_roof.jpg')
    try:
        from PIL import Image
        img = Image.new('RGB', (640, 480), color=(200, 180, 150))
        img.save(jpg_path, 'JPEG')
        paths.append(jpg_path)
    except ImportError:
        pass

    return paths


def make_docx_file():
    """DOCX file — uses python-docx if available."""
    try:
        from docx import Document
        path = os.path.join(OUT_DIR, 'proposal.docx')
        doc = Document()
        doc.add_heading('Roofing Proposal', 0)
        doc.add_heading('Customer: Sarah Johnson', level=1)
        doc.add_paragraph('Property: 142 Maple Street, Springfield, IL 62701')
        doc.add_paragraph('Date: 04/12/2024')
        doc.add_paragraph('')
        doc.add_heading('Scope of Work', level=1)
        doc.add_paragraph('1. Tear off existing 3-tab shingles (28 SQ)')
        doc.add_paragraph('2. Install GAF Timberline HDZ Charcoal shingles')
        doc.add_paragraph('3. Tiger Paw synthetic underlayment')
        doc.add_paragraph('4. New drip edge and ice & water shield')
        doc.add_heading('Investment', level=1)
        doc.add_paragraph('Total: $18,500.00')
        doc.add_paragraph('Warranty: 25 year manufacturer + 5 year workmanship')
        doc.save(path)
        return path
    except ImportError:
        print("  (python-docx not installed, skipping DOCX)")
        return None


if __name__ == '__main__':
    print(f"Generating test documents in {OUT_DIR}/...")
    paths = []

    paths.append(('searchable_pdf', make_searchable_pdf()))
    paths.append(('scanned_pdf', make_scanned_pdf()))
    paths.append(('mixed_pdf', make_mixed_pdf()))
    paths.append(('carrier_letter', make_carrier_letter_pdf()))
    paths.append(('inspection_report', make_inspection_report_pdf()))
    paths.append(('text_file', make_text_file()))
    paths.append(('csv_file', make_csv_file()))

    image_paths = make_image_files()
    for i, p in enumerate(image_paths):
        paths.append((f'image_{i}', p))

    docx_path = make_docx_file()
    if docx_path:
        paths.append(('docx_file', docx_path))

    print(f"\nGenerated {len(paths)} test documents:")
    for label, path in paths:
        size = os.path.getsize(path) if os.path.exists(path) else 0
        print(f"  {label:20s} → {path} ({size} bytes)")
